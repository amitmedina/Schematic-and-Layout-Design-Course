from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from xml.sax.saxutils import escape


def add_omml_equation(paragraph, linear: str) -> None:
    """Insert a Word equation object (OMML) containing the given linear math string."""
    eq_xml = (
        f'<m:oMathPara {nsdecls("m")}>'
        f"<m:oMath><m:r><m:t xml:space=\"preserve\">{escape(linear)}</m:t></m:r></m:oMath>"
        f"</m:oMathPara>"
    )
    paragraph._p.append(parse_xml(eq_xml))


def build_equations() -> list[dict[str, str]]:
    # Equations from LM5148 datasheet pages 36-39 (Eq.31–45).
    # Stored in a compact, linear form; Word will treat them as equation objects.
    return [
        {
            "num": "31",
            "title": "Buck inductance",
            "eq": "L = Vout*(Vin_nom - Vout)/(Vin_nom*Fsw*ΔILo)",
        },
        {
            "num": "32",
            "title": "Peak inductor current",
            "eq": "IL(pk) = Iout + (Vout*(Vin_max - Vout)/(Vin_max*L*Fsw))/2",
        },
        {
            "num": "33",
            "title": "Slope compensation cross-check (as shown)",
            "eq": "RO(sc) = Vout*Rs/(L*Fsw)",
        },
        {
            "num": "34",
            "title": "Current-sense resistor",
            "eq": "Rs = Vcs_th/(1.25*IL(pk))",
        },
        {
            "num": "35",
            "title": "Short-circuit peak inductor current",
            "eq": "IL_pk(sc) = (Vcs_th/Rs) + Vin_max*t_delay/L",
        },
        {
            "num": "36",
            "title": "Output capacitance for load-off overshoot",
            "eq": "Cout = L*Iout^2/( (Vout+ΔV)^2 - Vout^2 )",
        },
        {
            "num": "37",
            "title": "Output ripple (capacitive + ESR)",
            "eq": "ΔVout_pp ≈ sqrt( (ΔILo/(8*Fsw*Cout_eff))^2 + (ΔILo*RESR)^2 )",
        },
        {
            "num": "38",
            "title": "Output capacitor RMS ripple current",
            "eq": "Ico(rms) = ΔILo/sqrt(12)",
        },
        {
            "num": "39",
            "title": "Input capacitor RMS ripple current",
            "eq": "Icin(rms) = Iout*sqrt(D*(1-D))",
        },
        {
            "num": "40",
            "title": "Input capacitance (with ESR term)",
            "eq": "Choose Cin such that ΔVin_pp meets spec; one form uses ΔVin_cap ≈ Iout*D*(1-D)/(Fsw*Cin) with ESR ripple added",
        },
        {
            "num": "41",
            "title": "Frequency set resistor",
            "eq": "Fsw(kHz) = 10^6/(45*Rt(kΩ) + 53)",
        },
        {
            "num": "42",
            "title": "Feedback resistors",
            "eq": "Vout = Vref*(1 + Rtop/Rbottom)",
        },
        {
            "num": "43",
            "title": "RCOMP (per datasheet procedure)",
            "eq": "RCOMP = (Vout*Rs*Gm)/(2*π*fC*Cout_eff*Vref)",
        },
        {
            "num": "44",
            "title": "CCOMP zero placement",
            "eq": "CCOMP = 10/(2*π*fC*RCOMP)",
        },
        {
            "num": "45",
            "title": "CHF pole at ESR zero",
            "eq": "CHF = 1/(2*π*fESR*RCOMP) - Cbw",
        },
    ]


def main() -> int:
    parser = argparse.ArgumentParser(description="Export LM5148 datasheet equations (pages 36-39) to a Word .docx using Word equation objects.")
    parser.add_argument(
        "--out",
        type=str,
        default=str(Path.cwd() / "lm5148_equations_pages_36_39.docx"),
        help="Output .docx path",
    )
    args = parser.parse_args()

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("LM5148 Equations (Datasheet pages 36–39)", level=1)
    doc.add_paragraph(
        "Equations rewritten as Word equation objects (OMML). "
        "Symbols follow typical LM5148 datasheet notation (Vin, Vout, Fsw, ΔIL, etc.)."
    )

    equations = build_equations()
    for item in equations:
        doc.add_heading(f"Equation ({item['num']}): {item['title']}", level=2)
        p = doc.add_paragraph()
        add_omml_equation(p, item["eq"])

    doc.save(out_path)
    print(f"Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
